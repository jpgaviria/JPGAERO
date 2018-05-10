import xlsxwriter
import xlrd
import datetime
#import pandas as pd
#import gc
#import matplotlib.pyplot as plt
#import os
#import glob
#import math
#from docx import Document
import docx
#oldVDDpath = "VDD_generator_M145_V5_5_1_Delivery_JIRA Systems and Domains 3-19-2018"
newVDDpath = "Step1_VDD_2018-03-19-1452.xlsx"
SoftwareVersion  = "V5.5.0"
DocsVersion1 = "V5.5.0 Docs"
DocsVersion2 = "V5.5.1 Docs"
DocsVersion3 = "V5.5 Black Label Docs"
NewVDDdate  =   "March 15 2018"
#---------------------------------------------------------------
def write_new_HLS(HLS, TotalData, Links):
    now = datetime.datetime.now()
    VDDname = 'Step2_VDD_' + now.strftime("%Y-%m-%d-%H%M") +'.xlsx'
    DCPname = 'DCP_' + now.strftime("%Y-%m-%d-%H%M") +'.docx'

    wb = xlsxwriter.Workbook(VDDname)
    ws = wb.add_worksheet("HLS")
    ws2 = wb.add_worksheet("Total Data")
    ws3 = wb.add_worksheet("Link")
    document = docx.Document()
    
    position_format = wb.add_format()
    position_format.set_border()
    position_format.set_text_wrap()
    ws.set_column(0,0,10)
    ws.set_column(1,1,10)
    ws.set_column(2,2,100)
    ws.set_column(3,3,20)
    ws.set_column(4,4,20)

    ws2.set_column(0,0,15)
    ws2.set_column(1,1,60)
    ws2.set_column(2,2,15)
    ws2.set_column(3,3,15)
    ws2.set_column(4,4,15)
    ws2.set_column(5,5,15)
    ws2.set_column(6,6,20)
    ws2.set_column(7,7,15)
    ws2.set_column(8,8,15)
    ws2.set_column(9,9,15)
    ws2.set_column(10,10,30)
    ws2.set_column(11,11,120)
    for i in range(12,21,1):
        ws2.set_column(i,i,80)
    ws3.set_column(0,0,60)
    lines = len(HLS)
    rows = len(HLS[0])
    LinesTD = len(TotalData)
    columnsTD = len(TotalData[0])
    LinesLink = len(Links)
    #columnsLink = len(TotalData[0])
    for i in range (0,lines,1):
        for j in range (0, rows, 1):
            ws.write(i,j,str(HLS[i][j]),position_format)
    for i in range (0,LinesTD,1):
        for j in range (0, columnsTD, 1):
             ws2.write(i,j,str(TotalData[i][j]),position_format)
    #ws3.write(i,0,LinksTab.cell(i,0).value,position_format)
    #ws3.write(i,1,LinksTab.cell(i,1).value,position_format)
    for i in range (0,LinesLink,1):
        LinkCategory = str(Links[i])
        #LinkCategory = LinkCategory.replace('True(','')
        #LinkCategory = LinkCategory.replace(')','')
        LinkCategory = LinkCategory.replace('\\n','')
        LinkCategory = LinkCategory.replace('[','')
        LinkCategory = LinkCategory.replace(']','')
        LinkCategory = LinkCategory.replace("'","")
        ws3.write(i,0,LinkCategory,position_format)

    for i in range(0,lines,1):
        if str(HLS[i][3]) == '':
            document.add_heading(str(HLS[i][2]),level=1)
        else:
            document.add_paragraph(str(HLS[i][2]),style='List Bullet')
    document.save(DCPname)
    print (VDDname, ' Created\n',DCPname, ' Created')
#---------------------------------------------------------------
def find_duplicates(HLS):
    lines = len(HLS)
    print(lines)
    for i in range(0,lines,1):
        if str(HLS[i][3]) != '':
            for j in range((i+1),lines,1):
                if str(HLS[j][3]) != '':
                    if str(HLS[i][3]) == str(HLS[j][3]):
                        print(HLS[j][3], "is duplicate")
                        HLS[i][4] = "Duplicate on file"
                        HLS[j][4] = "Duplicate on file"
    return HLS
#---------------------------------------------------------------
def ReadTotalData():
    book = xlrd.open_workbook(newVDDpath)
    TotalDataTab = book.sheet_by_name('Total Data')    
    TDRows = len(TotalDataTab.col(7))
    TDColumns  =   len(TotalDataTab.row(1))
    TotalData = [['' for x in range(TDColumns)] for y in range(TDRows)]

    for i in range(0,TDRows,1):
        for j in range(0,TDColumns,1):
            TotalData[i][j] = str(TotalDataTab.cell(i,j).value)


    return TotalData
#---------------------------------------------------------------
def ReadLinks():
    book = xlrd.open_workbook(newVDDpath)
    LinksTab = book.sheet_by_name('Link')    
    LinksRows = len(LinksTab.col(0))
    LinksColumns  =   len(LinksTab.row(1))
    Links = [['' for x in range(LinksColumns)] for y in range(LinksRows)]

    for i in range(0,LinksRows,1):
        for j in range(0,LinksColumns,1):
            Links[i][j] = str(LinksTab.cell(i,j).value)            


    return Links
#---------------------------------------------------------------
def Update_HLS_from_TotalData(TotalData, Links):
    book = xlrd.open_workbook(newVDDpath)

    HLSTab = book.sheet_by_name('HLS')    
    HLSRows = len(HLSTab.col(3))
    HLSColumns  =   len(HLSTab.row(1))
    HLS = [['' for x in range(4)] for y in range(2)]

    for i in range(0,2,1):
        for j in range(0,4,1):
            HLS[i][j] = str(HLSTab.cell(i,j).value)            
   
    LinksRows = len(Links)
    TDRows = len(TotalData)


    for i in range(2,LinksRows,1):
        LinkCategory = str(Links[i])
        LinkCategory = LinkCategory.replace('True(','')
        LinkCategory = LinkCategory.replace(')','')
        LinkCategory = LinkCategory.replace('\\n','')
        LinkCategory = LinkCategory.replace('[','')
        LinkCategory = LinkCategory.replace(']','')
        LinkCategory = LinkCategory.replace("'","")
        HLS.append(['','',LinkCategory,''])
        for j in range(1, TDRows,1):
            TotalDataCategory = str(TotalData[j][6])
            TotalDataCategory = TotalDataCategory.replace('True(','')
            TotalDataCategory = TotalDataCategory.replace(')','')
            if TotalDataCategory==LinkCategory:
                HLS.append(['','',TotalData[j][1],TotalData[j][7]])


    return HLS
#-------------------------------------------------------------------
def main():
    TotalData = ReadTotalData()
    Links = ReadLinks()
    HLS = Update_HLS_from_TotalData(TotalData, Links)
    write_new_HLS(HLS, TotalData, Links)
#----------------------------------------------------------------------
if __name__ == "__main__":
    main()