import xlsxwriter
import xlrd
import datetime
#import pandas as pd
#import gc
#import matplotlib.pyplot as plt
#import os
#import glob
#import math
#from docx import Document
import docx
#oldVDDpath = "VDD_generator_M145_V5_5_1_Delivery_JIRA Systems and Domains 3-19-2018"
newVDDpath = "VDD_TotalData_4-11-2018.xlsx"
#SoftwareVersionSoftwareVersion  = "V5.5.0"
#DocsVersion1 = "V5.5.0 Docs"
#DocsVersion2 = "V5.5.1 Docs"
#DocsVersion3 = "V5.5 Black Label Docs"
#NewVDDdate  =   "March 15 2018"
#---------------------------------------------------------------
def write_new_HLS(HLS, TotalData, Links):
    now = datetime.datetime.now()
    VDDname = 'Step2_VDD_' + now.strftime("%Y-%m-%d-%H%M") +'.xlsx'
    DCPname = 'DCP_' + now.strftime("%Y-%m-%d-%H%M") +'.docx'

    wb = xlsxwriter.Workbook(VDDname)
    ws = wb.add_worksheet("HLS")
    ws2 = wb.add_worksheet("Total Data")
    ws3 = wb.add_worksheet("Link")
    document = docx.Document()
    
    position_format = wb.add_format()
    position_format.set_border()
    position_format.set_text_wrap()
    position_format.set_align("left")
    position_format.set_align("top")

    position_format_2 = wb.add_format()
    #date_form = wb.add_format({'num_format': 'mm/dd/yy'})
    #position_format_2.set_align("center")
    #position_format_2.set_align("vcenter")
    position_format_2.set_border()
    position_format_2.set_text_wrap()
    position_format_2.set_bold()
    position_format_2.set_bg_color("silver")
    position_format_2.set_font_color("Blue")
    position_format_2.set_align("left")
    position_format_2.set_align("top")

    
    
    date_format = wb.add_format({'num_format': 'mm/dd/yyyy'})
    date_format.set_border()
    date_format.set_align("align")
    date_format.set_align("valign")
    
    ws.set_column(0,0,10)
    ws.set_column(1,1,10)
    ws.set_column(2,2,100)
    ws.set_column(3,3,20)
    ws.set_column(4,4,20)

    ws2.set_column(0,0,15)
    ws2.set_column(1,1,60)
    ws2.set_column(2,2,15)
    ws2.set_column(3,3,15)
    ws2.set_column(4,4,15)
    ws2.set_column(5,5,15)
    ws2.set_column(6,6,20)
    ws2.set_column(7,7,15)
    ws2.set_column(8,8,15)
    ws2.set_column(9,9,15)
    ws2.set_column(10,10,30)
    ws2.set_column(11,11,120)
    for i in range(12,21,1):
        ws2.set_column(i,i,80)
    ws3.set_column(0,0,60)
    lines = len(HLS)
    rows = len(HLS[0])
    LinesTD = len(TotalData)
    columnsTD = len(TotalData[0])
    LinesLink = len(Links)
    #columnsLink = len(TotalData[0])
    for row in range(1,lines):
        i=0
        for char in HLS[row][3]:
            # if char == '\n':
            #     TotalData[row][7] = TotalData[row][7][:i]
            if (char == 'O' and HLS[row][3][i:i+3] == 'OMS' and i>3) and (HLS[row][3][i-1]!='\n'):
                HLS[row][3] = HLS[row][3][:i] + '\n' + HLS[row][3][i:]
                i+=2
                continue
            elif (char == 'I' and HLS[row][3][i:i+4] == 'IFIS' and i>3) and (HLS[row][3][i-1]!='\n'):
                HLS[row][3] = HLS[row][3][:i] + '\n' + HLS[row][3][i:]
                i+=2
                continue
            elif (char == 'E' and HLS[row][3][i:i+5] == 'EICAS' and i>3) and (HLS[row][3][i-1]!='\n'):
                HLS[row][3] = HLS[row][3][:i] + '\n' + HLS[row][3][i:]
                i+=2
                continue            
            elif (char == 'F' and HLS[row][3][i:i+4] == 'FDSA' and i>3) and (HLS[row][3][i-1]!='\n'):
                HLS[row][3] = HLS[row][3][:i] + '\n' + HLS[row][3][i:]
                i+=2
                continue            
            else:
                i+=1
    
    for i in range (0,lines,1):
        for j in range (0, rows, 1):
            if (j==3) and (str(HLS[i][j]) == '') and (str(HLS[i][j-1]) != ''):
                ws.write(i,(j-1),str(HLS[i][j-1]),position_format_2)
                ws.write(i,j,str(HLS[i][j]),position_format_2)
            else:    
                ws.write(i,j,str(HLS[i][j]),position_format)
    for row in range(1,LinesTD):
        i=0
        for char in TotalData[row][7]:
            # if char == '\n':
            #     TotalData[row][7] = TotalData[row][7][:i]
            if (char == 'O' and TotalData[row][7][i:i+3] == 'OMS' and i>3) and (TotalData[row][7][i-1]!='\n'):
                TotalData[row][7] = TotalData[row][7][:i] + '\n' + TotalData[row][7][i:]
            elif (char == 'I' and TotalData[row][7][i:i+4] == 'IFIS' and i>3) and (TotalData[row][7][i-1]!='\n'):
                TotalData[row][7] = TotalData[row][7][:i] + '\n' + TotalData[row][7][i:]
            elif (char == 'E' and TotalData[row][7][i:i+5] == 'EICAS' and i>3) and (TotalData[row][7][i-1]!='\n'):
                TotalData[row][7] = TotalData[row][7][:i] + '\n' + TotalData[row][7][i:]
            elif (char == 'F' and TotalData[row][7][i:i+4] == 'FDSA' and i>3) and (TotalData[row][7][i-1]!='\n'):
                TotalData[row][7] = TotalData[row][7][:i] + '\n' + TotalData[row][7][i:]
            i+=1

    for i in range (0,LinesTD,1):
        for j in range (0, columnsTD, 1):
            if j==2 or j==3:
                ws2.write(i,j,TotalData[i][j],date_format)
            else:
                ws2.write(i,j,str(TotalData[i][j]),position_format)
    #ws3.write(i,0,LinksTab.cell(i,0).value,position_format)
    #ws3.write(i,1,LinksTab.cell(i,1).value,position_format)
    for i in range (0,LinesLink,1):
        LinkCategory = str(Links[i])
        #LinkCategory = LinkCategory.replace('True(','')
        #LinkCategory = LinkCategory.replace(')','')
        LinkCategory = LinkCategory.replace('\\n','')
        LinkCategory = LinkCategory.replace('[','')
        LinkCategory = LinkCategory.replace(']','')
        LinkCategory = LinkCategory.replace("'","")
        ws3.write(i,0,LinkCategory,position_format)

    for i in range(0,lines,1):
        if str(HLS[i][3]) == '':
            document.add_heading(str(HLS[i][2]),level=1)
        else:
            document.add_paragraph(str(HLS[i][2]),style='List Bullet')
    document.save(DCPname)
    print (VDDname, ' Created\n',DCPname, ' Created')
#---------------------------------------------------------------
def find_duplicates(HLS):
    lines = len(HLS)
    print(lines)
    for i in range(0,lines,1):
        if str(HLS[i][3]) != '':
            for j in range((i+1),lines,1):
                if str(HLS[j][3]) != '':
                    if str(HLS[i][3]) == str(HLS[j][3]):
                        print(HLS[j][3], "is duplicate")
                        HLS[i][4] = "Duplicate on file"
                        HLS[j][4] = "Duplicate on file"
    return HLS
#---------------------------------------------------------------
def ReadTotalData():
    book = xlrd.open_workbook(newVDDpath)
    TotalDataTab = book.sheet_by_name('Total Data')    
    TDRows = len(TotalDataTab.col(7))
    TDColumns  =   len(TotalDataTab.row(1))
    TotalData = [['' for x in range(TDColumns)] for y in range(TDRows)]
    for i in range(0,TDRows,1):
        for j in range(0,TDColumns,1):
            cellstring = isinstance(TotalDataTab.cell(i,j).value,str)
            if(j == 2 or j==3 or j == 4 or j==5) and i>0 and (cellstring==False):
                TotalData[i][j] = xlrd.xldate.xldate_as_datetime(TotalDataTab.cell(i,j).value, book.datemode)
            else:
                TotalData[i][j] = str(TotalDataTab.cell(i,j).value)
    #search Total data and replace o with real bullets
    for i in range(0,TDRows,1):
        string1 = str(TotalData[i][1])
        #string1 = string1.replace('\u2022','')
        string1 = string1.replace('\no','\n\u2022\u0009')
        string1 = string1.replace('\n o','\n\u0009\u2022\u0009')
        string1 = string1.replace('\n  o','\n  \u0009\u2022\u0009')
        TotalData[i][1] = str(string1)

    return TotalData
#---------------------------------------------------------------
def ReadLinks():
    book = xlrd.open_workbook(newVDDpath)
    LinksTab = book.sheet_by_name('Link')    
    LinksRows = len(LinksTab.col(0))
    LinksColumns  =   len(LinksTab.row(1))
    Links = [['' for x in range(LinksColumns)] for y in range(LinksRows)]

    for i in range(0,LinksRows,1):
        for j in range(0,LinksColumns,1):
            Links[i][j] = str(LinksTab.cell(i,j).value)            


    return Links
#---------------------------------------------------------------
def Update_HLS_from_TotalData(TotalData, Links):
    book = xlrd.open_workbook(newVDDpath)

    HLSTab = book.sheet_by_name('HLS')    
    #HLSRows = len(HLSTab.col(3))
    #HLSColumns  =   len(HLSTab.row(1))
    HLS = [['' for x in range(4)] for y in range(2)]

    for i in range(0,2,1):
        for j in range(0,4,1):
            HLS[i][j] = str(HLSTab.cell(i,j).value)            
   
    LinksRows = len(Links)
    TDRows = len(TotalData)


    for i in range(2,LinksRows,1):
        LinkCategory = str(Links[i])
        LinkCategory = LinkCategory.replace('True(','')
        LinkCategory = LinkCategory.replace(')','')
        LinkCategory = LinkCategory.replace('\\n','')
        LinkCategory = LinkCategory.replace('[','')
        LinkCategory = LinkCategory.replace(']','')
        LinkCategory = LinkCategory.replace("'","")
        HLS.append(['','','',''])
        HLS.append(['','',LinkCategory,''])
        for j in range(1, TDRows,1):
            TotalDataCategory = str(TotalData[j][6])
            TotalDataCategory = TotalDataCategory.replace('True(','')
            TotalDataCategory = TotalDataCategory.replace(')','')
            TotalDataCategory = TotalDataCategory.replace('\\n','')
            TotalDataCategory = TotalDataCategory.replace('\n','')
            TotalDataCategory = TotalDataCategory.replace("'","")
            if TotalDataCategory==LinkCategory:
                HLS.append(['','',TotalData[j][1],TotalData[j][7]])


    return HLS
#-------------------------------------------------------------------
def main():
    TotalData = ReadTotalData()
    Links = ReadLinks()
    HLS = Update_HLS_from_TotalData(TotalData, Links)
    write_new_HLS(HLS, TotalData, Links)
#----------------------------------------------------------------------
if __name__ == "__main__":
    main()