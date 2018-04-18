import xlsxwriter
import xlrd
import datetime
#import pandas as pd
#import gc
#import matplotlib.pyplot as plt
#import os
#import glob
#import math
#from docx import Document
import docx
from jira import JIRA
import os
oldVDDpath = "VDD_generator_M145_V5_5_1_Delivery_JIRA Systems and Domains 4-17-2018.xlsx"
newVDDpath = "VDD_generator_M185_V5_5_1_Cert - 23 March 2018.xlsx"
SoftwareVersion  = "V5.5.0"
SoftwareVersion2  = "V5.5.1"
DocsVersion1 = "V5.5.0 Docs"
DocsVersion2 = "V5.5.1 Docs"
DocsVersion3 = "V5.5 Black Label Docs"
NewVDDdate  =   datetime.datetime(2018,3,23)
os.chdir('C:\\Users\jpgaviri\iCloudDrive\Personal\Python\VDD_Management')
#----------------------------------------------------------------------------
def Get_Update_from_Jira(TotalData):
    print ("Fetching updates form Jira, make sure you are on the RCI network and set the password on the script")

    NumRows = len(TotalData)
    CRs = ['' for x in range(NumRows)]
    NumColumn = len(TotalData[0])
    print(NumColumn)
    for row in range(0,NumRows):
        TotalData[row].insert(0,'Jira')
    NumColumn = len(TotalData[0])
    print(NumColumn)

    for row in range(0,NumRows):
        CRs[row] = str(TotalData[row][8])
    # cleanup list of CRs so that only 1 CR per
    for row in range(1,len(CRs)):
        i=0
        for char in CRs[row]:
            if char == '\n':
                CRs[row] = CRs[row][:i]
            elif char == 'O' and CRs[row][i:i+3] == 'OMS' and i>3:
                CRs[row] = CRs[row][:i]
            elif char == 'I' and CRs[row][i:i+4] == 'IFIS' and i>3:
                CRs[row] = CRs[row][:i]
            elif char == 'E' and CRs[row][i:i+5] == 'EICAS' and i>3:
                CRs[row] = CRs[row][:i]
            elif char == 'F' and CRs[row][i:i+4] == 'FDSA' and i>3:
                CRs[row] = CRs[row][:i]
            i+=1
    options = {'server': 'http://alm.rockwellcollins.com/issues/'}
    jira = JIRA(options, basic_auth=('jpgaviri','Canada18*'))#Use RCI password

    i=0
    for CR in CRs:
        if CR != '' and CR != 'CRID' and CR != 'GLOBALA-XXXX' and CR != 'Globala-XXX' and CR != 'GLOBALA-TBD':
            Wp_type, summary, description, function, fix_version, affected_projects,labels, \
            Operational_impact,Plain_English,Criticality = '','','','','','','','','',''
            if CR[:6] == 'GLOBAL':
                issue = jira.issue(str(CR))
                print (CR)
                Wp_type, summary, description, function, fix_version, affected_projects,labels, \
                Operational_impact,Plain_English,Criticality = '','','','','','','','','',''
                for field in issue.raw['fields']:
                    if field == 'customfield_12909':
                        if issue.raw['fields']['customfield_12909'] != None:
                            Wp_type = issue.raw['fields']['customfield_12909']['value']
                    if field == 'summary':
                        if issue.raw['fields']['summary'] != None:
                            summary = issue.raw['fields']['summary']
                    if field == 'description':
                        if issue.raw['fields']['description'] != None:
                            description = issue.raw['fields']['description']
                    if field == 'customfield_19307':
                        if issue.raw['fields']['customfield_19307'] != None:
                            function = issue.raw['fields']['customfield_19307']['value']
                    if field == 'fixVersions':
                        if issue.raw['fields']['fixVersions'] != []:
                            fix_version = issue.raw['fields']['fixVersions'][0]['name']
                    if field == 'customfield_13705':
                        if issue.raw['fields']['customfield_13705'] != None:
                            affected_projects = issue.raw['fields']['customfield_13705']
                    if field == 'labels':
                        if issue.raw['fields']['labels'] != None:
                            labels = issue.raw['fields']['labels']
                    if field == 'customfield_12906':
                        if issue.raw['fields']['customfield_12906'] != None:
                            Operational_impact = issue.raw['fields']['customfield_12906']
                    if field == 'customfield_18527':
                        if issue.raw['fields']['customfield_18527'] != None:
                            Plain_English = issue.raw['fields']['customfield_18527']
                    if field == 'customfield_13514':
                        if issue.raw['fields']['customfield_13514'] != []:
                            Criticality = issue.raw['fields']['customfield_13514'][0]
                #last_update = issue.raw['fields']['customfield_19063']
                print ("this seems to have worked")
                label = ''
                if labels != '':
                    for l in labels: 
                        label = label + ', ' + str(l)
                label = str(label)

                TotalData[i][0] = 'CR NUmber: ' + str(issue.key) + '\n' \
                                    'Work Package Type: ' + str(Wp_type) + '\n' \
                                    + 'Fix Version: '+ str(fix_version) +'\n' \
                                    + 'Labels: '+ str(label) + '\n' \
                                    + 'Summary: ' + str(summary) + '\n' \
                                    + 'Operational Impact: '+ str(Operational_impact) +'\n' \
                                    + 'English Description: '+ str(Plain_English) +'\n' \
                                    + 'Description: ' + str(description) +'\n'
                #AllM145Issues[i][1].value = issue.key
                i+=1
            else:
                i+=1
        else:
            i+=1
    return TotalData


#-----------------------------------------------------------------------------
def add_CRs_from_new_VDD(TotalData):
    book = xlrd.open_workbook(newVDDpath)

    PlannedTab = book.sheet_by_name('Planned')
    #SWversionTab = book.sheet_by_name(SoftwareVersion)
    #Docs1Tab = book.sheet_by_name(DocsVersion1)
    #Docs2Tab = book.sheet_by_name(DocsVersion2)
    #Docs3Tab = book.sheet_by_name(DocsVersion3)
    PlannedColumns = len(PlannedTab.col(2))
    PlannedRows = len(PlannedTab.row(0))
    Planned = [[0 for x in range(PlannedRows)] for y in range(PlannedColumns)] 
    for i in range(0,PlannedColumns,1):
        for j in range(0,PlannedRows,1):
            cellstring = isinstance(PlannedTab.cell(i,j).value,str)
            if(j == 2 or j==3 or j == 4 or j==5) and i>0 and (cellstring==False):
                Planned[i][j] = xlrd.xldate.xldate_as_datetime(PlannedTab.cell(i,j).value, book.datemode)
            else:
                Planned[i][j] = str(PlannedTab.cell(i,j).value)
    CRfound = False
    while CRfound == False:
        for i in range(1,PlannedColumns,1):
            CRfound = False
            CRNumberPlanned = Planned[i][2]
            CRNumberPlanned = CRNumberPlanned.replace(' ','')
            CRNumberPlanned = CRNumberPlanned.replace('\n','')
            for j in range(1,len(TotalData),1):
                CRnumberTD = TotalData[j][7]
                CRnumberTD = CRnumberTD.replace(' ','')
                CRnumberTD = CRnumberTD.replace('\n','')
                if CRnumberTD == CRNumberPlanned:
                    CRfound = True
                    break
            if CRfound == False:
                #TotalData.append(Planned[i])
                TotalData.append(['','','','','','','','','','','','','','','','','','','','',''])
                #TotalData[j+1].extend(['','','','','','',''])
                TotalData[j+1][0] = ''
                TotalData[j+1][1] = 'New CR from planning tab'
                TotalData[j+1][2] = NewVDDdate
                TotalData[j+1][3] = NewVDDdate
                TotalData[j+1][4] = ''
                TotalData[j+1][5] = ''
                #Linked
                TotalData[j+1][6] = ''
                TotalData[j+1][7] = CRNumberPlanned
                TotalData[j+1][8] = Planned[i][0]
                TotalData[j+1][9] = Planned[i][1]
                for k in range(10,len(Planned[0]),1):
                    TotalData[j+1][k] = Planned[i][k-7]
                CRfound = True
                print(CRNumberPlanned, " from Planned tab does not exist in Total data")
    #-------------------------SW CRs---------------------------------------------        
    SWTab = book.sheet_by_name(SoftwareVersion)
#    Docs1Tab = book.sheet_by_name(DocsVersion1)
#    Docs2Tab = book.sheet_by_name(DocsVersion2)
#    Docs3Tab = book.sheet_by_name(DocsVersion3)
    SWColumns = len(SWTab.col(2))
    SWRows = len(SWTab.row(0))
    SW = [[0 for x in range(SWRows)] for y in range(SWColumns)] 
    for i in range(0,SWColumns,1):
        for j in range(0,SWRows,1):
            cellstring = isinstance(SWTab.cell(i,j).value,str)
            if(j == 2 or j==3 or j == 4 or j==5) and i>0 and (cellstring==False):
                SW[i][j] = xlrd.xldate.xldate_as_datetime(SWTab.cell(i,j).value, book.datemode)
            else:
                SW[i][j] = str(SWTab.cell(i,j).value)
    CRfound = False
    while CRfound == False:
        for i in range(1,SWColumns,1):
            CRfound = False
            CRNumberSW = SW[i][2]
            CRNumberSW = CRNumberSW.replace(' ','')
            CRNumberSW = CRNumberSW.replace('\n','')
            for j in range(1,len(TotalData),1):
                CRnumberTD = TotalData[j][7]
                CRnumberTD = CRnumberTD.replace(' ','')
                CRnumberTD = CRnumberTD.replace('\n','')
                if CRnumberTD == CRNumberSW:
                    CRfound = True
                    break
            if CRfound == False:
                #TotalData.append(Planned[i])
                TotalData.append(['','','','','','','','','','','','','','','','','','','','',''])
                #TotalData[j+1].extend(['','','','','','',''])
                TotalData[j+1][0] = ''
                TotalData[j+1][1] = 'New CR from '+ SoftwareVersion +' SW tab'
                TotalData[j+1][2] = NewVDDdate
                TotalData[j+1][3] = NewVDDdate
                TotalData[j+1][4] = ''
                TotalData[j+1][5] = ''
                #Linked
                TotalData[j+1][6] = ''
                TotalData[j+1][7] = CRNumberSW
                TotalData[j+1][8] = SW[i][0]
                TotalData[j+1][9] = SW[i][1]
                for k in range(10,18,1):
                    TotalData[j+1][k] = SW[i][k-7]
                TotalData[j+1][18] = 'N/A'
                TotalData[j+1][19] = SW[i][11]
                TotalData[j+1][20] = SW[i][12]
                CRfound = True
                print(CRNumberSW, " from SW tab does not exist in Total data")
    #-------------------------SW CRs---------------------------------------------        
    SWTab2 = book.sheet_by_name(SoftwareVersion2)
#    Docs1Tab = book.sheet_by_name(DocsVersion1)
#    Docs2Tab = book.sheet_by_name(DocsVersion2)
#    Docs3Tab = book.sheet_by_name(DocsVersion3)
    SWColumns = len(SWTab2.col(2))
    SWRows = len(SWTab2.row(0))
    SW2 = [[0 for x in range(SWRows)] for y in range(SWColumns)] 
    for i in range(0,SWColumns,1):
        for j in range(0,SWRows,1):
            cellstring = isinstance(SWTab2.cell(i,j).value,str)
            if(j == 2 or j==3 or j == 4 or j==5) and i>0 and (cellstring==False):
                SW2[i][j] = xlrd.xldate.xldate_as_datetime(SWTab2.cell(i,j).value, book.datemode)
            else:
                SW2[i][j] = str(SWTab2.cell(i,j).value)
    CRfound = False
    while CRfound == False:
        for i in range(1,SWColumns,1):
            CRfound = False
            CRNumberSW = SW2[i][2]
            CRNumberSW = CRNumberSW.replace(' ','')
            CRNumberSW = CRNumberSW.replace('\n','')
            for j in range(1,len(TotalData),1):
                CRnumberTD = TotalData[j][7]
                CRnumberTD = CRnumberTD.replace(' ','')
                CRnumberTD = CRnumberTD.replace('\n','')
                if CRnumberTD == CRNumberSW:
                    CRfound = True
                    break
            if CRfound == False:
                #TotalData.append(Planned[i])
                TotalData.append(['','','','','','','','','','','','','','','','','','','','',''])
                #TotalData[j+1].extend(['','','','','','',''])
                TotalData[j+1][0] = ''
                TotalData[j+1][1] = 'New CR from '+ SoftwareVersion2 +' SW tab'
                TotalData[j+1][2] = NewVDDdate
                TotalData[j+1][3] = NewVDDdate
                TotalData[j+1][4] = ''
                TotalData[j+1][5] = ''
                #Linked
                TotalData[j+1][6] = ''
                TotalData[j+1][7] = CRNumberSW
                TotalData[j+1][8] = SW2[i][0]
                TotalData[j+1][9] = SW2[i][1]
                for k in range(10,18,1):
                    TotalData[j+1][k] = SW2[i][k-7]
                TotalData[j+1][18] = 'N/A'
                TotalData[j+1][19] = SW2[i][11]
                TotalData[j+1][20] = SW2[i][12]
                CRfound = True
                print(CRNumberSW, " from SW2 tab does not exist in Total data")

    #-------------------------DOCS 1--------------------------------------------- 
#    SWTab = book.sheet_by_name(SoftwareVersion)
    Docs1Tab = book.sheet_by_name(DocsVersion1)
#    Docs2Tab = book.sheet_by_name(DocsVersion2)
#    Docs3Tab = book.sheet_by_name(DocsVersion3)
    Docs1Columns = len(Docs1Tab.col(2))
    Docs1Rows = len(Docs1Tab.row(0))
    Docs1 = [[0 for x in range(Docs1Rows)] for y in range(Docs1Columns)] 
    for i in range(0,Docs1Columns,1):
        for j in range(0,Docs1Rows,1):
            cellstring = isinstance(Docs1Tab.cell(i,j).value,str)
            if(j == 2 or j==3 or j == 4 or j==5) and i>0 and (cellstring==False):
                Docs1[i][j] = xlrd.xldate.xldate_as_datetime(Docs1Tab.cell(i,j).value, book.datemode)
            else:
                Docs1[i][j] = str(Docs1Tab.cell(i,j).value)
    CRfound = False
    while CRfound == False:
        for i in range(1,Docs1Columns,1):
            CRfound = False
            CRNumberDocs1 = Docs1[i][2]
            CRNumberDocs1 = CRNumberDocs1.replace(' ','')
            CRNumberDocs1 = CRNumberDocs1.replace('\n','')
            for j in range(1,len(TotalData),1):
                CRnumberTD = TotalData[j][7]
                CRnumberTD = CRnumberTD.replace(' ','')
                CRnumberTD = CRnumberTD.replace('\n','')
                if CRnumberTD == CRNumberDocs1:
                    CRfound = True
                    break
            if CRfound == False:
                #TotalData.append(Planned[i])
                TotalData.append(['','','','','','','','','','','','','','','','','','','','',''])
                #TotalData[j+1].extend(['','','','','','',''])
                TotalData[j+1][0] = ''
                TotalData[j+1][1] = 'New CR from '+ DocsVersion1 +' SW tab'
                TotalData[j+1][2] = NewVDDdate
                TotalData[j+1][3] = NewVDDdate
                TotalData[j+1][4] = ''
                TotalData[j+1][5] = ''
                #Linked
                TotalData[j+1][6] = ''
                TotalData[j+1][7] = CRNumberDocs1
                TotalData[j+1][8] = Docs1[i][0]
                TotalData[j+1][9] = Docs1[i][1]
                for k in range(10,18,1):
                    TotalData[j+1][k] = Docs1[i][k-7]
                TotalData[j+1][18] = 'N/A'
                TotalData[j+1][19] = Docs1[i][11]
                TotalData[j+1][20] = Docs1[i][12]
                CRfound = True
                print(CRNumberDocs1, " from Docs1 tab does not exist in Total data")
       
    #-------------------------DOCS 2---------------------------------------------
#    SWTab = book.sheet_by_name(SoftwareVersion)
#    Docs1Tab = book.sheet_by_name(DocsVersion1)
    Docs2Tab = book.sheet_by_name(DocsVersion2)
#    Docs3Tab = book.sheet_by_name(DocsVersion3)
    Docs2Columns = len(Docs2Tab.col(2))
    Docs2Rows = len(Docs2Tab.row(0))
    Docs2 = [[0 for x in range(Docs2Rows)] for y in range(Docs2Columns)] 
    for i in range(0,Docs2Columns,1):
        for j in range(0,Docs2Rows,1):
            cellstring = isinstance(Docs2Tab.cell(i,j).value,str)
            if(j == 2 or j==3 or j == 4 or j==5) and i>0 and (cellstring==False):
                Docs2[i][j] = xlrd.xldate.xldate_as_datetime(Docs2Tab.cell(i,j).value, book.datemode)
            else:
                Docs2[i][j] = str(Docs2Tab.cell(i,j).value)
    CRfound = False
    while CRfound == False:
        for i in range(1,Docs2Columns,1):
            CRfound = False
            CRNumberDocs2 = Docs2[i][2]
            CRNumberDocs2 = CRNumberDocs2.replace(' ','')
            CRNumberDocs2 = CRNumberDocs2.replace('\n','')
            for j in range(1,len(TotalData),1):
                CRnumberTD = TotalData[j][7]
                CRnumberTD = CRnumberTD.replace(' ','')
                CRnumberTD = CRnumberTD.replace('\n','')
                if CRnumberTD == CRNumberDocs2:
                    CRfound = True
                    break
            if CRfound == False:
                #TotalData.append(Planned[i])
                TotalData.append(['','','','','','','','','','','','','','','','','','','','',''])
                #TotalData[j+1].extend(['','','','','','',''])
                TotalData[j+1][0] = ''
                TotalData[j+1][1] = 'New CR from '+ DocsVersion2 +' SW tab'
                TotalData[j+1][2] = NewVDDdate
                TotalData[j+1][3] = NewVDDdate
                TotalData[j+1][4] = ''
                TotalData[j+1][5] = ''
                #Linked
                TotalData[j+1][6] = ''
                TotalData[j+1][7] = CRNumberDocs2
                TotalData[j+1][8] = Docs2[i][0]
                TotalData[j+1][9] = Docs2[i][1]
                for k in range(10,18,1):
                    TotalData[j+1][k] = Docs2[i][k-7]
                TotalData[j+1][18] = 'N/A'
                TotalData[j+1][19] = Docs2[i][11]
                TotalData[j+1][20] = Docs2[i][12]
                CRfound = True
                print(CRNumberDocs2, " from Docs2 tab does not exist in Total data")
        
    #-------------------------DOCS 3---------------------------------------------        
#    SWTab = book.sheet_by_name(SoftwareVersion)
#    Docs1Tab = book.sheet_by_name(DocsVersion1)
#    Docs2Tab = book.sheet_by_name(DocsVersion2)
    Docs3Tab = book.sheet_by_name(DocsVersion3)
    Docs3Columns = len(Docs3Tab.col(2))
    Docs3Rows = len(Docs3Tab.row(0))
    Docs3 = [[0 for x in range(Docs3Rows)] for y in range(Docs3Columns)] 
    for i in range(0,Docs3Columns,1):
        for j in range(0,Docs3Rows,1):
            cellstring = isinstance(Docs3Tab.cell(i,j).value,str)
            if(j == 2 or j==3 or j == 4 or j==5) and i>0 and (cellstring==False):
                Docs3[i][j] = xlrd.xldate.xldate_as_datetime(Docs3Tab.cell(i,j).value, book.datemode)
            else:
                Docs3[i][j] = str(Docs3Tab.cell(i,j).value)
    CRfound = False
    while CRfound == False:
        for i in range(1,Docs3Columns,1):
            CRfound = False
            CRNumberDocs3 = Docs3[i][2]
            CRNumberDocs3 = CRNumberDocs3.replace(' ','')
            CRNumberDocs3 = CRNumberDocs3.replace('\n','')
            for j in range(1,len(TotalData),1):
                CRnumberTD = TotalData[j][7]
                CRnumberTD = CRnumberTD.replace(' ','')
                CRnumberTD = CRnumberTD.replace('\n','')
                if CRnumberTD == CRNumberDocs3:
                    CRfound = True
                    break
            if CRfound == False:
                #TotalData.append(Planned[i])
                TotalData.append(['','','','','','','','','','','','','','','','','','','','',''])
                #TotalData[j+1].extend(['','','','','','',''])
                TotalData[j+1][0] = ''
                TotalData[j+1][1] = 'New CR from '+ DocsVersion3 +' SW tab'
                TotalData[j+1][2] = NewVDDdate
                TotalData[j+1][3] = NewVDDdate
                TotalData[j+1][4] = ''
                TotalData[j+1][5] = ''
                #Linked
                TotalData[j+1][6] = ''
                TotalData[j+1][7] = CRNumberDocs3
                TotalData[j+1][8] = Docs3[i][0]
                TotalData[j+1][9] = Docs3[i][1]
                for k in range(10,18,1):
                    TotalData[j+1][k] = Docs3[i][k-7]
                TotalData[j+1][18] = 'N/A'
                TotalData[j+1][19] = Docs3[i][11]
                TotalData[j+1][20] = Docs3[i][12]
                CRfound = True
                print(CRNumberDocs3, " from Docs3 tab does not exist in Total data")
  
    return TotalData
#-----------------------------------------------------------------------------
def create_TotalData_from_HLS(HLS):
    HLSLines = len(HLS)
    #HLSColumns = len(HLS[0])
    TotalData = [['' for x in range(21)] for y in range(HLSLines)] 
    for i in range(0,HLSLines,1):
        if str(HLS[i][3]) != '':
            TotalData[i][7] = HLS[i][3]
            TotalData[i][1] = HLS[i][2]
            #CRnumHLS = HLS[i][3]
            #CRnumHLS = CRnumHLS.replace(' ','')
            #CRnumHLS = CRnumHLS.replace('\n','')
            #TotalData[i][7] = HLS[i][3]
            #for j in range(0,HLSLines,1):
            #    if str(TotalData[j][7])!='':
            #        CRnumTD = TotalData[j][7]
            #        CRnumTD = CRnumTD.replace(' ','')
            #        CRnumTD = CRnumTD.replace('\n','')
            #        if CRnumHLS == CRnumTD:
            #            TotalData[j][1] = HLS[i][2]
    for i in range(0,HLSLines,1):
        string1 = str(TotalData[i][1])
        #string1 = string1.replace('\u2022 ','')
        string1 = string1.replace('\no','\n\u2022')
        string1 = string1.replace('\n o','\n\u0009\u2022')
        string1 = string1.replace('\n  o','\n  \u0009\u2022')
        TotalData[i][1] = str(string1)

    return TotalData


#----------------------------------------------------------------------------- 
def update_TotalData_from_PrevTotalData(TotalData):
    """
    Open read and paste contents into a list an Excel file
    """
    book = xlrd.open_workbook(oldVDDpath)
    #write = xlsxwriter.Workbook(oldVDDpath)
    # print number of sheets
    X = book.sheet_by_name('Total Data')
    PrevTDcolums = len(X.col(4))
    PrevTDrows = len(X.row(1))
    PrevTotalData = [[0 for x in range(PrevTDrows)] for y in range(PrevTDcolums)] 
    #print (PrevTDcolums, PrevTDrows)
    for i in range(0,PrevTDcolums,1):
        for j in range(0,PrevTDrows,1):
            cellstring = isinstance(X.cell(i,j).value,str)
            if(j == 2 or j==3 or j == 4 or j==5) and i>0 and (cellstring==False):
                PrevTotalData[i][j] = xlrd.xldate.xldate_as_datetime(X.cell(i,j).value, book.datemode)
            else:
                PrevTotalData[i][j] = str(X.cell(i,j).value)
    for i in range(0,PrevTDcolums,1):
        #isdate = isinstance(PrevTotalData[i][3],datetime.datetime)
        if ' ' in str(PrevTotalData[i][7]):
            word = str(PrevTotalData[i][7])
            word = word.replace(' ','')
            word = word.replace('\n','')
            PrevTotalData[i][7] = word
    for i in range(0,PrevTDcolums,1):
        if ' ' in str(PrevTotalData[i][7]):
            print('CR with spaces')
    TDcolums = len(TotalData)
    #TDrows = len(TotalData[0])
    TotalData[0] = PrevTotalData[0]
    # fill out all other colums other than CR ID on new total data
    for i in range(1,TDcolums, 1):
        CRnumTD = TotalData[i][7]
        CRnumTD = CRnumTD.replace(' ','')
        CRnumTD = CRnumTD.replace('\n','')
        for j in range(1,PrevTDcolums,1):
            prevCRnumTD = PrevTotalData[j][7]
            prevCRnumTD = prevCRnumTD.replace(' ','')
            prevCRnumTD = prevCRnumTD.replace('\n','')
            if CRnumTD == prevCRnumTD:
                TotalData[i][0] = PrevTotalData[j][0]
                TotalData[i][2] = PrevTotalData[j][2]
                TotalData[i][3] = NewVDDdate
                for k in range(4,7,1):
                    TotalData[i][k] = PrevTotalData[j][k]
                for k in range(8,len(PrevTotalData[0]),1):
                    TotalData[i][k] = PrevTotalData[j][k]
    # Add items from previous total data not in the new total data
    for i in range(1,PrevTDcolums, 1):
        prevCRnumTD = PrevTotalData[i][7]
        prevCRnumTD = prevCRnumTD.replace(' ','')
        prevCRnumTD = prevCRnumTD.replace('\n','')
        found = False
        while found == False:
            found = False
            for j in range(1,TDcolums,1):
                CRnumTD = TotalData[j][7]
                CRnumTD = CRnumTD.replace(' ','')
                CRnumTD = CRnumTD.replace('\n','')
                if CRnumTD == prevCRnumTD:
                    found = True
            if found == False:
                TotalData.append(PrevTotalData[i])
                found = True
                print(prevCRnumTD, " does not exist in HLS")


    return TotalData
#---------------------------------------------------------------
def write_new_HLS_and_TotalData(HLS, TotalData):
    book = xlrd.open_workbook(oldVDDpath)
    now = datetime.datetime.now()
    VDDname = 'Step1_VDD_' + now.strftime("%Y-%m-%d-%H%M") +'.xlsx'

    LinksTab = book.sheet_by_name('Link')
    wb = xlsxwriter.Workbook(VDDname)
    ws = wb.add_worksheet("HLS")
    ws2 = wb.add_worksheet("Total Data")
    ws3 = wb.add_worksheet("Link")
    document = docx.Document()
    
    position_format = wb.add_format()
    #date_form = wb.add_format({'num_format': 'mm/dd/yy'})
    position_format.set_align("left")
    position_format.set_align("top")
    position_format.set_border()
    position_format.set_text_wrap()
    #position_format.set_bold()
    #position_format.set_bg_color("silver")
    date_format = wb.add_format({'num_format': 'mm/dd/yyyy'})
    date_format.set_border()
    date_format.set_align("align")
    date_format.set_align("valign")
    ws.set_column(0,0,10)
    ws.set_column(1,1,10)
    ws.set_column(2,2,40)
    ws.set_column(3,3,20)
    ws.set_column(4,4,20)
    
    # ws2.set_column(0,0,15)
    # ws2.set_column(1,1,60)
    # ws2.set_column(2,2,15)
    # ws2.set_column(3,3,15)
    # ws2.set_column(4,4,15)
    # ws2.set_column(5,5,15)
    # ws2.set_column(6,6,20)
    # ws2.set_column(7,7,15)
    # ws2.set_column(8,8,15)
    # ws2.set_column(9,9,15)
    # ws2.set_column(10,10,30)
    # ws2.set_column(11,11,120)

    ws2.set_column(0,0,100)
    ws2.set_column(1,1,15)
    ws2.set_column(2,2,60)
    ws2.set_column(3,3,15)
    ws2.set_column(4,4,15)
    ws2.set_column(5,5,15)
    ws2.set_column(6,6,15)
    ws2.set_column(7,7,20)
    ws2.set_column(8,8,15)
    ws2.set_column(9,9,15)
    ws2.set_column(10,10,15)
    ws2.set_column(11,11,30)
    ws2.set_column(12,12,120)

    for i in range(13,21,1):
        ws2.set_column(i,i,80)
    ws3.set_column(0,0,60)
    lines = len(HLS)
    rows = len(HLS[0])
    LinesTD = len(TotalData)
    columnsTD = len(TotalData[0])
    LinesLink = len(LinksTab.col(0))
    #columnsLink = len(TotalData[0])
    for i in range (0,lines,1):
        for j in range (0, rows, 1):
            ws.write(i,j,str(HLS[i][j]),position_format)
    for i in range (0,LinesTD,1):
        for j in range (0, columnsTD, 1):
            if j== 2 or j==3:
                ws2.write(i,j,TotalData[i][j],date_format)
            else:
                ws2.write(i,j,str(TotalData[i][j]),position_format)
    #ws3.write(i,0,LinksTab.cell(i,0).value,position_format)
    #ws3.write(i,1,LinksTab.cell(i,1).value,position_format)
    for i in range (0,LinesLink,1):
        ws3.write(i,0,str(LinksTab.cell(i,0).value),position_format)

    for i in range(0,lines,1):
        if str(HLS[i][3]) == '':
            document.add_heading(str(HLS[i][2]),level=1)
        else:
            document.add_paragraph(str(HLS[i][2]),style='List Bullet')
    document.save('DCP.docx')
    print(VDDname, ' Created')
#---------------------------------------------------------------
def find_duplicates(HLS):
    lines = len(HLS)
    print(lines)
    for i in range(0,lines,1):
        if str(HLS[i][3]) != '':
            for j in range((i+1),lines,1):
                if str(HLS[j][3]) != '':
                    str1 = str(HLS[i][3])
                    str1 = str1.upper()
                    str2 = str(HLS[j][3])
                    str2 = str2.upper()
                    if str1 == str2:
                        print(HLS[j][3], "is duplicate")
                        HLS[i][4] = "Duplicate on file"
                        HLS[j][4] = "Duplicate on file"
    return HLS
#----------------------------------------------------------------------
def open_file():
    """
    Open readand paste contents into a list an Excel file
    """
    book = xlrd.open_workbook(oldVDDpath)
    #write = xlsxwriter.Workbook(oldVDDpath)
    # print number of sheets
    X = book.sheet_by_name('HLS')
    lines = len(X.col(3))
    rows = len(X.row(1))
    HLS = [[0 for x in range(rows)] for y in range(lines)] 
    #print (lines, rows)
    for i in range(0,lines,1):
        for j in range(0,rows,1):
            HLS[i][j] = str(X.cell(i,j).value)
    for i in range(0,lines,1):
        if ' ' in str(HLS[i][3]):
            word = str(HLS[i][3])
            word = word.replace(' ','')
            word = word.replace('\n','')
            HLS[i][3] = word
    for i in range(0,lines,1):
        if ' ' in str(HLS[i][3]):
            print('CR with spaces')
    return HLS

#---------------------------------------------------------------
def delete_Empty_Lines(TotalData):
    
    #TDrows = len(TotalData[0])
    dupe = False
    endloop = False
    while endloop == False:
        dupe = False
        TDcolums = len(TotalData)
        for i in range(1,TDcolums,1):
            if str(TotalData[i][7])=='' and dupe ==False:
                TotalData.pop(i)
                dupe = True
                break
        if i == (TDcolums-1):
            endloop = True

    return TotalData 
#-------------------------------------------------------------------
def main():
    HLS = open_file()
    HLS = find_duplicates(HLS)
    TotalData = create_TotalData_from_HLS(HLS)
    TotalData = update_TotalData_from_PrevTotalData(TotalData)
    TotalData = delete_Empty_Lines(TotalData)
    TotalData = add_CRs_from_new_VDD(TotalData)
    TotalData = Get_Update_from_Jira(TotalData)
    write_new_HLS_and_TotalData(HLS,TotalData)
#----------------------------------------------------------------------
if __name__ == "__main__":
    main()